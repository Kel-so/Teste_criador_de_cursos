import streamlit as st
import google.generativeai as genai
import json
from PyPDF2 import PdfReader
from docx import Document
import io

# ==========================================
# 1. Configuração e Variáveis Fixas
# ==========================================
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

modelo_estrutura = genai.GenerativeModel(
    "gemini-3.1-flash-lite-preview", 
    generation_config={"response_mime_type": "application/json"}
)
modelo_texto = genai.GenerativeModel("gemini-3.1-flash-lite-preview")

# Base do Saber Completa
BASE_DO_SABER_TEXTO = """
BASE do SABER
Estrutura Oficial para Produção de Cursos - Saber Gestão

Introdução
No Saber Gestão, acreditamos que educação de verdade transforma comportamento e gera resultado. Para garantir que nossos cursos entreguem essa transformação, estamos desenvolvendo o Algoritmo do SABER - uma metodologia completa e validada que será a base de todos os nossos produtos educacionais.
Enquanto essa metodologia está em fase final de validação, estruturamos a BASE do SABER - uma estrutura didática clara, prática e alinhada com nossos princípios, que orienta a produção de cursos com consistência, aplicabilidade e qualidade.

A BASE do SABER tem como objetivo:
- Padronizar a criação de cursos, mesmo com temas, formatos e instrutores diferentes;
- Garantir coerência didática e fluidez na experiência do aluno;
- Assegurar que todo curso entregue valor real: Despertando Curiosidade, Fundamentando o Conteúdo, Promovendo aplicação Prática e Conduzindo o aluno a Ação.

Essa base não é uma simples sequência de etapas. Ela é uma orientação estratégica e andragógica, que garante consistência sem engessar a criatividade. Os cursos podem ser técnicos, comportamentais, normativos ou estratégicos e essa estrutura se adapta a todos, desde que os objetivos de cada módulo sejam respeitados.

Fase 1 - Despertar a Curiosidade
Objetivo: Gerar interesse genuíno, criar conexão emocional e apresentar o "porque" do curso.
Aulas obrigatórias:
Aula 1 - Introdução e Apresentação do Curso: Apresente o tema, objetivos e benefícios diretos para o aluno.
Aula 2 - Provocação / Pergunta Reflexiva: Levante uma pergunta poderosa que instigue pensamento crítico e curiosidade.
Aula 3 - Observação Inicial. Escolher entre: Desmistificação: "Do que não se trata esse tema" (evita distorções comuns), ou Case real ou história de aplicação prática, preferencialmente impactante e direto.
Aula 4 - Conclusão do Módulo: Autodiagnóstico sobre o conhecimento ou prática atual - SaberForm ou estrutura compatível dentro da ferramenta; Fechamento com reforço da importância do aprendizado.

Fase 2 - Fundamentação Técnica
Objetivo: Explicar com profundidade e clareza técnica o "o que é" do tema, oferecendo múltiplas formas de assimilação.
Diretrizes:
- Sem número fixo de aulas: devem ser incluídas quantas forem necessárias para garantir compreensão sólida;
- Cada aula deve tratar um conceito específico, com exemplos práticos e linguagem acessível.

Fase 3 - Aplicação e Atividades
Objetivo: Ensinar como aplicar o que foi aprendido e estimular a prática ativa.
Elementos obrigatórios (sem limitação de número de aulas):
- Exemplos de aplicação prática: Casos reais, simulações, passo a passo, erros comuns, boas práticas;
- Estudo de caso ou simulação: Situação real ou hipotética onde o conhecimento pode ser testado;
- Atividade prática orientada: Desafio, roteiro, checklist, canvas ou formulário guiado;
- Quiz de reforço ou autoavaliação prática.

Fase 4 - Ação e Reflexão Final
Objetivo: Induzir o aluno a ação prática ou a reflexão estratégica, com foco em transformação real.
Aulas obrigatórias:
Aula 1 - Ação recomendada no cotidiano e Reflexão final.
Aula 2 - Fechamento do curso e Formulário de Satisfação.

Modelo de Aula - Estrutura Detalhada (Novo Saber Play)
Ordem e Função de Cada Bloco OBRIGATÓRIO:
0 | Título Impactante: Captar a atenção imediatamente (Curioso, direto, provocativo)
1 | Frase de Abertura: Criar expectativa e gerar conexão (Pode ser emocional, racional ou provocativa, 1-2 linhas)
2 | Bloco de Vídeo: Entregar explicação base e conexão humana (Duração: 3 a 7 minutos)
3 | Bloco de Texto Guiado: Reforçar e estruturar o conteúdo aprendido (Frases curtas, listas, microexemplos)
4 | Bloco de Aplicação: Traduzir teoria em ação prática (Checklist, passo a passo, roteiro)
5 | Reflexão ou Escrita Ativa: Estimular personalização do aprendizado (Campo de anotação ou instrução para registro)
6 | Quiz ou Avaliação Rápida (opcional): Verificar fixação e reforçar aprendizagem (3 a 5 perguntas contextualizadas - sempre com 4 alternativas)
7 | Bloco Complementar (opcional): Expandir visão para alunos protagonistas (Links, PDFs, materiais extras)

Fluxo Narrativo da Aula: Desperta -> Conecta -> Ensina -> Estrutura -> Aplica -> Faz pensar -> Reforça -> Expande.

Regras Visuais e de Escrita:
- Evite parágrafos longos (máximo 3 linhas por bloco)
- Use destaque visual (negrito, bullets, emojis leves)
- Títulos dos blocos devem ser curtos e funcionais
- Evite jargões desnecessários: clareza vence o tecnicismo
- Evitar criar aulas com quantidade exagerada de blocos
- Evitar criar aulas em que o conteúdo aplicado é excessivamente extensivo
- Evitar aulas diferentes com blocos redundantes
- A organização dos blocos deve ser visualmente agradável
"""

# ==========================================
# 2. Funções de Suporte e Templates
# ==========================================
def extrair_texto(arquivo):
    if arquivo is None:
        return ""
    if arquivo.type == "application/pdf":
        leitor = PdfReader(arquivo)
        texto = ""
        for pagina in leitor.pages:
            texto += pagina.extract_text()
        return texto
    elif arquivo.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(io.BytesIO(arquivo.read()))
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return arquivo.read().decode("utf-8")

def gerar_docx(texto):
    doc = Document()
    for linha in texto.split('\n'):
        doc.add_paragraph(linha)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

TEMPLATES = {
    "Roteiro de Aula": """Atue como um instrutor técnico especializado e conteudista sênior da Saber Gestão. 
Crie o roteiro de gravação DETALHADO para a aula: {aula}.

REGRA ABSOLUTA: Siga rigorosamente a estrutura da Base do Saber.

Use formatação rica (negrito, tópicos) e emojis. Seu retorno DEVE conter OBRIGATORIAMENTE os seguintes blocos:
- 🎯 Título Impactante (Curioso e direto)
- 💬 Frase de Abertura (Para gerar conexão)
- 🎥 Bloco de Vídeo: 
   ATENÇÃO: Este é o roteiro exato da fala do professor (estilo teleprompter). 
   TOM DE VOZ: Conversacional, humano, empático e direto. Parágrafos curtos (1 a 3 frases). Conecte com a vida real e traga a parte técnica mastigada.
   FORMATO DE EXIBIÇÃO: {formato}. (Adapte o roteiro se for animado, gravado no estúdio ou misto).
- 📝 Bloco de Texto Guiado (Resumo em tópicos curtos, máximo de 3 linhas por tópico)
- 🛠️ Bloco de Aplicação (Passo a passo prático para aplicar no dia a dia)
- 🤔 Bloco de Reflexão (Pergunta estratégica para o aluno pensar)
- ❓ Quiz (1 ou 2 questões com 4 alternativas focadas na prática do conteúdo)
- 🎯 Resultado da Aula (Lista objetiva do que o aluno atingiu ao final desta aula)

Base de Produção Nativa:
{base_saber}

Contexto Técnico da Norma/Tema: {norma}
Diretrizes do Projeto Pedagógico (Se houver): {pedagogico}""",

    "Roteiro de Materiais": "Liste todos os equipamentos, EPIs e ferramentas visuais necessárias para demonstrar na prática os conceitos da aula: {aula}. Base: {norma}.",
    "Prompts IA (Imagens)": "Crie 3 prompts detalhados em inglês para gerar imagens de apoio visual para a aula: {aula}. As imagens devem ter estilo realista, iluminação de estúdio, proporção 16:9 e focar no tema central da aula. Apenas entregue os prompts em texto puro."
}

# ==========================================
# 3. Interface Visual e Estados
# ==========================================
st.set_page_config(page_title="Studio Saber - IA", layout="wide", page_icon="⚙️")
st.title("⚙️ Gerador de Estrutura e Roteiros")

if "estrutura_curso" not in st.session_state:
    st.session_state.estrutura_curso = None
if "roteiros_gerados" not in st.session_state:
    st.session_state.roteiros_gerados = {}

with st.sidebar:
    st.header("📂 Arquivos Base (Opcional)")
    st.caption("A Base do Saber já está integrada nativamente.")
    file_norma = st.file_uploader("Norma (PDF/Word)", type=["pdf", "docx", "txt"])
    file_pedagogico = st.file_uploader("Projeto Pedagógico", type=["pdf", "docx", "txt"])
    
    st.divider()
    
    st.header("📝 Preenchimento Manual")
    st.caption("Use os campos abaixo caso não possua os arquivos prontos.")
    manual_tema = st.text_input("Tema Central do Curso", placeholder="Ex: NR-35 ou Liderança")
    manual_norma = st.text_input("Norma/Requisito", placeholder="Ex: ISO 9001, CLT...")
    manual_publico = st.text_input("Público-Alvo", placeholder="Ex: Operadores, Supervisores")
    manual_objetivo = st.text_area("Objetivo Principal", placeholder="O que o aluno deve ser capaz de fazer no final?", height=100)

    tema_manual = ""
    if manual_tema or manual_norma or manual_publico or manual_objetivo:
        tema_manual = f"Tema Central: {manual_tema}\nNorma de Referência: {manual_norma}\nPúblico-Alvo: {manual_publico}\nObjetivo: {manual_objetivo}"

norma_texto = extrair_texto(file_norma) or tema_manual
pedagogico_texto = extrair_texto(file_pedagogico)

aba_estrutura, aba_roteiros = st.tabs(["🏗️ 1. Planejar Estrutura", "📝 2. Produção de Roteiros"])

# ==========================================
# ABA 1: Configuração do Curso
# ==========================================
with aba_estrutura:
    st.markdown("### Configuração do Projeto")
    
    col_horas, col_mods = st.columns(2)
    with col_horas:
        carga_horaria = st.number_input("Carga Horária (Horas)", min_value=1, value=16, help="Define o volume de aulas nas fases de Fundamentação e Aplicação.")
    with col_mods:
        num_modulos = st.number_input("Quantidade de Módulos", min_value=2, value=4)
        
    st.markdown("### Estrutura Macro dos Módulos")
    modulos_config = []
    
    for i in range(int(num_modulos)):
        cm1, cm2, cm3 = st.columns([2, 2, 1])
        with cm1:
            nome_m = st.text_input("Nome do Módulo", value=f"Módulo {i+1}", key=f"nm_{i}")
        with cm2:
            fase_padrao = 0 if i == 0 else (3 if i == int(num_modulos) - 1 else 1)
            fase_m = st.selectbox("Fase do Saber", ["Fase 1 - Despertar a Curiosidade", "Fase 2 - Fundamentação Técnica", "Fase 3 - Aplicação Prática", "Fase 4 - Ação e Reflexão Final"], index=fase_padrao, key=f"fs_{i}")
        with cm3:
            formato_m = st.selectbox("Formato", ["Gravado", "Animado", "Misto"], key=f"fm_{i}")
            
        modulos_config.append({"nome": f"{nome_m} ({formato_m})", "fase": fase_m, "formato": formato_m})

    st.divider()
    
    pode_gerar = bool(norma_texto.strip())
    if not pode_gerar:
        st.warning("⚠️ Faça o upload de uma Norma ou preencha os campos Manuais na barra lateral para liberar a geração.")
        
    if st.button("🚀 Gerar Estrutura Completa", disabled=not pode_gerar, use_container_width=True):
        prompt_json = f"""
        Você é um arquiteto de cursos conteudista sênior da Saber Gestão, especialista em normas regulamentadoras e treinamentos corporativos.
        Crie a estrutura do curso solicitada e retorne ESTRITAMENTE em formato JSON.
        
        REGRAS PEDAGÓGICAS NATIVAS (BASE DO SABER):
        {BASE_DO_SABER_TEXTO}
        
        TEMA/NORMA BASE: {norma_texto}
        PROJETO PEDAGÓGICO DE APOIO: {pedagogico_texto}
        
        CONFIGURAÇÃO EXIGIDA:
        - Carga Horária Total: {carga_horaria} horas.
        - Estrutura de Módulos: {json.dumps(modulos_config, ensure_ascii=False)}
        
        REGRAS ABSOLUTAS E SEPARAÇÃO DE FASES:
        1. MÓDULOS DE FASE 1 (Despertar): DEVEM ter exatamente as 4 aulas prescritas na Base do Saber. O campo 'topicos_norma' pode ser "N/A".
        2. MÓDULOS DE FASE 4 (Ação/Reflexão): DEVEM ter exatamente as 2 aulas prescritas na Base do Saber. O campo 'topicos_norma' pode ser "N/A".
        3. MÓDULOS DE FASE 2 (Fundamentação Técnica - Módulo 2): É OBRIGATÓRIO extrair e cobrir 100% do CONTEÚDO PROGRAMÁTICO TEÓRICO exigido pela norma. Aloque esses tópicos rigorosamente no campo 'topicos_norma'. NENHUM item teórico exigido pela norma pode ficar de fora.
        4. MÓDULOS DE FASE 3 (Aplicação Prática - Módulo 3): Este módulo NÃO PODE ter aulas teóricas. Crie EXCLUSIVAMENTE aulas de práticas, estudos de caso, simulações, uso de checklists, passo a passo e resolução de problemas reais. O campo 'topicos_norma' deve referenciar a prática do conteúdo abordado.
        5. DOSAGEM PELA CARGA HORÁRIA: Use a carga horária de {carga_horaria} horas para ditar a fragmentação nas fases 2 e 3. Se a carga for alta, crie mais aulas aprofundadas. Se for baixa, agrupe os tópicos.
        
        Retorne o JSON neste formato exato (substitua os valores com a sua estruturação):
        {{
            "modulos": [
                {{
                    "nome": "Módulo 2 - Fundamentação",
                    "fase": "Fase 2 - Fundamentação Técnica",
                    "formato": "Gravado",
                    "aulas": [
                        {{
                            "titulo": "Aula 2.1 - Definições e Reconhecimento",
                            "topicos_norma": "NR-33, Itens 33.1 e 33.2"
                        }}
                    ]
                }}
            ]
        }}
        """
        
        with st.spinner("Mapeando a norma e desenhando a estrutura pedagógica..."):
            try:
                resposta = modelo_estrutura.generate_content(prompt_json)
                st.session_state.estrutura_curso = json.loads(resposta.text)
                st.success("✅ Estrutura gerada com sucesso! Tópicos da norma mapeados. Vá para a aba 'Produção de Roteiros'.")
            except Exception as e:
                st.error("Erro ao processar a estrutura JSON. Verifique os dados fornecidos e tente novamente.")

# ==========================================
# ABA 2: Edição e Geração de Conteúdo
# ==========================================
with aba_roteiros:
    if st.session_state.estrutura_curso:
        st.markdown("### 🗂️ Estrutura e Roteiros")
        
        for idx_mod, modulo in enumerate(st.session_state.estrutura_curso["modulos"]):
            formato_txt = modulo.get("formato", "Misto")
            with st.expander(f"{modulo['nome']} [{formato_txt}]", expanded=True):
                
                novo_nome_mod = st.text_input("Nome do Módulo", value=modulo["nome"], key=f"mod_{idx_mod}")
                st.session_state.estrutura_curso["modulos"][idx_mod]["nome"] = novo_nome_mod
                st.divider() 
                
                for idx_aula, aula_obj in enumerate(modulo["aulas"]):
                    chave_roteiro = f"{idx_mod}_{idx_aula}"
                    texto_gerado = st.session_state.roteiros_gerados.get(chave_roteiro)
                    
                    # Compatibilidade (caso o JSON venha com strings ou dicionários)
                    if isinstance(aula_obj, str):
                        titulo_atual = aula_obj
                        topicos_atual = ""
                    else:
                        titulo_atual = aula_obj.get("titulo", "")
                        topicos_atual = aula_obj.get("topicos_norma", "")
                    
                    col_aula, col_tipo, col_btn, col_ler, col_baixar = st.columns([5, 3, 2, 1, 1], vertical_alignment="center")
                    
                    with col_aula:
                        nova_aula = st.text_input("Aula", value=titulo_atual, key=f"aula_tit_{chave_roteiro}", label_visibility="collapsed")
                        
                        # Exibe os tópicos da norma embaixo do nome da aula
                        if topicos_atual and topicos_atual not in ["N/A", "N/A.", ""]:
                            st.caption(f"📌 **Norma/Treinamento:** {topicos_atual}")
                            
                        # Salva de volta no estado
                        if isinstance(aula_obj, str):
                            st.session_state.estrutura_curso["modulos"][idx_mod]["aulas"][idx_aula] = nova_aula
                        else:
                            st.session_state.estrutura_curso["modulos"][idx_mod]["aulas"][idx_aula]["titulo"] = nova_aula
                            
                    with col_tipo:
                        tipo_geracao = st.selectbox("Template", list(TEMPLATES.keys()), key=f"tipo_{chave_roteiro}", label_visibility="collapsed")
                    
                    with col_btn:
                        gerar_clicado = st.button("✨ Gerar IA", key=f"btn_{chave_roteiro}", use_container_width=True)
                    
                    with col_ler:
                        if texto_gerado:
                            with st.popover("👁️", use_container_width=True):
                                st.markdown(texto_gerado)
                    
                    with col_baixar:
                        if texto_gerado:
                            docx_file = gerar_docx(texto_gerado)
                            st.download_button(
                                label="📥",
                                data=docx_file,
                                file_name=f"Roteiro_{nova_aula[:15].strip()}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"down_{chave_roteiro}",
                                use_container_width=True
                            )
                            
                    if gerar_clicado:
                        # Junta o título da aula e os tópicos obrigatórios para forçar a IA no prompt
                        aula_contexto = nova_aula
                        if topicos_atual and topicos_atual not in ["N/A", "N/A.", ""]:
                            aula_contexto += f" | Foco Obrigatório na Norma: {topicos_atual}"

                        formato_aula = modulo.get("formato", "Gravado")
                        prompt_final = TEMPLATES[tipo_geracao].format(
                            aula=aula_contexto, 
                            pedagogico=pedagogico_texto, 
                            norma=norma_texto,
                            formato=formato_aula,
                            base_saber=BASE_DO_SABER_TEXTO
                        )
                        with st.spinner(f"Escrevendo {tipo_geracao}..."):
                            resultado = modelo_texto.generate_content(prompt_final)
                            st.session_state.roteiros_gerados[chave_roteiro] = resultado.text
                            st.rerun() 
                            
    else:
        st.info("A estrutura do curso aparecerá aqui após ser gerada na aba anterior.")
